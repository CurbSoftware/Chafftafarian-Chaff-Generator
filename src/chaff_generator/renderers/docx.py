"""DOCX renderer — python-docx over ProseDocument (spec sections 66, 74).

Word bytes are not byte-deterministic across runs (OOXML carries zip
metadata); the manifest hash recorded at generation time is authoritative,
and tests assert semantic equality via python-docx parsing instead.
"""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx import Document as new_document
from docx.document import Document as DocxDocument
from docx.shared import Pt

from chaff_generator.content import builders
from chaff_generator.core.hashing import hash_file
from chaff_generator.renderers.base import RendererCapabilities, RenderResult

if TYPE_CHECKING:
    from pathlib import Path

    from chaff_generator.content.context import RenderContext
    from chaff_generator.renderers.base import Renderer
    from chaff_generator.renderers.documents import (
        Block,
        Paragraph,
        ProseDocument,
        SemanticDocument,
    )

CAPABILITIES = RendererCapabilities(
    extension="docx",
    supports_exact_size=False,
    supports_target_size=True,
    supports_streaming=False,
    semantic_document="prose",
    size_category="document",
)

#: The document body aims at this fraction of ``desired_size``; OOXML
#: framing makes up the rest (approximate-size format).
_BODY_FRACTION = 0.92

#: Keep the in-memory docx bounded: python-docx materializes the whole tree.
_MAX_BODY_CHARS = 60 << 20


def ensure_prose_volume(document: ProseDocument, context: RenderContext, target_chars: int) -> None:
    """Append appendix filler sections until the body reaches ``target_chars``.

    Shared by the docx and pdf renderers, which serialize the whole document
    at once rather than streaming (approximate-size formats).
    """
    target_chars = min(target_chars, _MAX_BODY_CHARS)
    from chaff_generator.renderers.documents import Paragraph, Section

    current = sum(
        len(block_text(block)) for section in document.sections for block in section.blocks
    )
    appendix: list[Paragraph] = []
    appendix_size = 0
    while current + appendix_size < target_chars:
        filler = builders.filler_paragraph(context, sentences=4)
        appendix.append(Paragraph(filler))
        appendix_size += len(filler) + 1
    if appendix:
        document.sections.append(Section(heading="Appendix", blocks=list(appendix)))


def block_text(block: Block) -> str:
    """Plain-text projection of one block (for size estimation)."""
    from chaff_generator.renderers.documents import (
        BulletList,
        Heading,
        NumberedList,
        Paragraph,
        Quote,
        Table,
    )

    if isinstance(block, Heading):
        return block.text
    if isinstance(block, (Paragraph, Quote)):
        return block.text
    if isinstance(block, (BulletList, NumberedList)):
        return " ".join(block.items)
    if isinstance(block, Table):
        return " ".join(block.columns) + " " + " ".join(" ".join(row) for row in block.rows)
    return ""


class DocxRenderer:
    id = "docx"
    capabilities = CAPABILITIES

    def render(
        self,
        document: SemanticDocument | None,
        destination: Path,
        context: RenderContext,
    ) -> RenderResult:
        if document is None:
            from chaff_generator.renderers.markdown import _fallback_document

            document = _fallback_document(context)
        from chaff_generator.renderers.documents import ProseDocument

        assert isinstance(document, ProseDocument)
        ensure_prose_volume(document, context, int(context.desired_size * _BODY_FRACTION))

        doc = new_document()
        self._apply_core_properties(doc, document)
        style = doc.styles["Normal"]
        style.font.size = Pt(11)

        doc.add_heading(document.title, level=0)
        for section in document.sections:
            if section.heading:
                doc.add_heading(section.heading, level=1)
            for block in section.blocks:
                self._write_block(doc, block)
        doc.save(str(destination))

        size = destination.stat().st_size
        return RenderResult(
            path=destination,
            size=size,
            renderer_id=self.id,
            template_id=context.template_id,
            sha256=hash_file(destination),
        )

    # ---------------------------------------------------------------- internals

    def _apply_core_properties(self, doc: DocxDocument, document: ProseDocument) -> None:
        from datetime import datetime
        from datetime import time as dt_time

        props = doc.core_properties
        props.title = document.title
        props.author = document.author
        props.last_modified_by = document.author
        stamp = datetime.combine(document.created_at, dt_time(9, 0))
        props.created = stamp
        props.modified = stamp
        props.comments = "Generated synthetic document"

    def _write_block(self, doc: DocxDocument, block: Block) -> None:
        from chaff_generator.renderers.documents import (
            BulletList,
            Heading,
            NumberedList,
            PageBreak,
            Paragraph,
            Quote,
            Table,
        )

        if isinstance(block, Heading):
            doc.add_heading(block.text, level=block.level)
        elif isinstance(block, Paragraph):
            doc.add_paragraph(block.text)
        elif isinstance(block, Quote):
            doc.add_paragraph(block.text, style="Intense Quote")
        elif isinstance(block, BulletList):
            for item in block.items:
                doc.add_paragraph(item, style="List Bullet")
        elif isinstance(block, NumberedList):
            for item in block.items:
                doc.add_paragraph(item, style="List Number")
        elif isinstance(block, Table):
            table = doc.add_table(rows=1 + len(block.rows), cols=len(block.columns))
            table.style = "Light Grid Accent 1"
            for column, name in enumerate(block.columns):
                table.rows[0].cells[column].text = str(name)
            for row_index, row in enumerate(block.rows, start=1):
                for column, value in enumerate(row):
                    table.rows[row_index].cells[column].text = str(value)
        elif isinstance(block, PageBreak):
            doc.add_page_break()


def get_renderer(renderer_id: str) -> Renderer:
    if renderer_id != "docx":
        raise ValueError(f"docx module cannot serve renderer id {renderer_id!r}")
    return DocxRenderer()
